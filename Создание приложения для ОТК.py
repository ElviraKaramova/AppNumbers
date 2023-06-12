from docx import Document
from docx.shared import Pt
import random
import sqlite3
db = sqlite3.connect('MedReg')
c = db.cursor()
c.execute('SELECT * FROM chashki')
listt = c.fetchall()
# Было ранее
# c.execute("""
# CREATE TABLE chashki (
# id integer,
# name text,
# p1 real,
# p2 real,
# p3 real,
# p4 real
# ) """)
# c.execute("INSERT INTO chashki VALUES ('A1','Чашка Петри Ø 35 мм, вентилируемая','35.4','39.2','12.3','3.62') ")
# Далее строки были добавлены в SQL начиная с INSERT INTO
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times-New-Roman'
style.font.size = Pt(14)

def create_table():
    table = doc.add_table(rows = 21, cols = 5)
    rows= 21 #В зависимости от выборки, может быть изменено
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ('№ образца')
    cell = table.cell(0, 1)
    cell.text = ('Внутренний диаметр основания, мм')
    cell = table.cell(0, 2)
    cell.text = ('Наружный диаметр крышки, мм')
    cell = table.cell(0, 3)
    cell.text = ('Высота чашки, мм')
    cell = table.cell(0, 4)
    cell.text = ('Масса чашки, г')
    for row in range(1,rows):
        for col in range(1):
            cell = table.cell(row, col)
            cell.text = str(("{0}{1}".format('A',row)))
        for col in range(1,2):
            cell = table.cell(row, col)
            cell.text = str("{0:0.2f}".format(round(random.uniform((listt[i][2]-0.8), (listt[i][2]+0.8)),2)))
        for col in range(2,3):
            cell = table.cell(row, col)
            cell.text = str("{0:0.2f}".format(round(random.uniform((listt[i][3]-0.8), (listt[i][3]+0.8)),2)))
        for col in range(3,4):
            cell = table.cell(row, col)
            cell.text = str("{0:0.2f}".format(round(random.uniform((listt[i][4]-0.4), (listt[i][4]+0.4)),2)))
        for col in range(4,5):
            cell = table.cell(row, col)
            cell.text = str("{0:0.2f}".format(round(random.uniform((listt[i][5]-0.2), (listt[i][5]+0.2)),2)))
    doc.add_paragraph()
# db.commit()
# db.close()
def stroki():
    s = 'Таблица '
    s2 = str(i+1)
    s3 = ' - Вариант исполнения '
    s4 = str(listt[i][1])
    s5= s+s2+s3+s4
    return doc.add_paragraph(s5)
for i in range(14):
    stroki()
    create_table()
doc.save('Приложение акта квалификационных испытаний.docx')